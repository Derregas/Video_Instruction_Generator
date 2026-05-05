import os
import json
import logging
from abc import ABC, abstractmethod
from typing import Type, Dict, Literal
from werkzeug.datastructures import FileStorage

from fpdf import FPDF
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def docs_size(docs: list[FileStorage]) -> int:
    total_size = 0
    for doc in docs:
        doc.seek(0, os.SEEK_END)    # Перемещаем указатель в конец файла
        size = doc.tell()           # Получаем позицию указателя в байтах - размер файла
        doc.seek(0)                 # возвращаем указатель в начало
        total_size += size
    return total_size

def docs_save(docs: list[FileStorage], request_temp_dir: str) -> list[str]:
    document_paths = []
    for doc in docs:
        doc_path = os.path.join(request_temp_dir, f"doc_{doc.filename}")
        doc.save(doc_path)
        document_paths.append(doc_path)
    return document_paths

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Manual PC Assembly', 0, 1, 'C')

class BaseDocument(ABC):
    base_fonts_path = os.path.join(os.path.dirname(__file__), 'static', 'fonts')
    @abstractmethod
    def create(self, content: str, filename: str) -> None:
        pass
    @staticmethod
    def _to_json(content: str):
        return json.loads(content)
    @staticmethod
    def format_time(seconds: float) -> str:
        m = int(seconds // 60)
        s = int(seconds % 60)
        return f"{m:02}:{s:02}"

class DocxDocument(BaseDocument):
    # TODO: Проще не заморачиваться со стилями, а создать шаблон документа с
    # готовыми стилями заголовков, так и редактировать проще будет
    @staticmethod
    def set_style_font(style, font_name):
        # Стандартная установка
        style.font.name = font_name
        # Глубокая установка через XML для всех типов символов (ascii, кириллица и т.д.)
        rFonts = style.element.rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    @staticmethod
    def set_style(document):
        if 'Title' in document.styles:
            t_style = document.styles['Title']
            DocxDocument.set_style_font(t_style, 'Times New Roman')
            t_style.font.size = Pt(18)
            t_style.font.bold = True
            t_style.font.color.rgb = None # Сброс цвета
            # Убираем рамку программно
            pPr = t_style.element.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                pPr.remove(pBdr)
        if 'Heading 1' in document.styles:
            h1_style = document.styles['Heading 1']
            DocxDocument.set_style_font(h1_style, 'Times New Roman')
            h1_style.font.size = Pt(14)
            h1_style.font.bold = True
            h1_style.font.color.rgb = None
        if 'Normal' in document.styles:
            normal = document.styles['Normal']
            DocxDocument.set_style_font(normal, 'Times New Roman')
            normal.font.size = Pt(14)

    def create(self, content: str, filename: str) -> None:
        doc = Document()
        json_content = self._to_json(content)
        # Настраиваем основные стили
        self.set_style(doc)
        # --- 1. НАСТРОЙКА ОТСТУПОВ СТРАНИЦЫ ---
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

        title = doc.add_heading('Инструкция по сборке', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Times New Roman'

        for step in json_content['instructions']:
            # Добавляем заголовок шага
            heading = doc.add_heading(step['title'], level=1)
            for run in heading.runs:
                run.font.name = 'Times New Roman'

            # Описание
            p = doc.add_paragraph(step['description'])
            p.paragraph_format.first_line_indent = Cm(1.25) # Красная строка
            p.paragraph_format.space_after = Pt(10)         # Отступ после абзаца
            
            # ВСТАВКА КАРТИНКИ
            img_path = os.path.abspath(step['best_image_id'])
            
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Cm(12))
                    # Центрируем картинку (она считается как отдельный параграф)
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.error(f"Ошибка вставки картинки {img_path}: {e}")
                    doc.add_paragraph(f"[Ошибка изображения: {os.path.basename(img_path)}]")
            else:
                logger.warning(f"Файл не найден: {img_path}")

            # Таймкоды
            time_text = f"Таймкод: {self.format_time(float(step['start_time']))} - {self.format_time(float(step['end_time']))}"
            caption = doc.add_paragraph(time_text)
            caption.italic = True # type: ignore
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(filename)
        
class PdfDocument(BaseDocument):

    def _setup_fonts(self, pdf):
        """Вспомогательный метод для регистрации шрифтов"""
        font_styles: list[tuple[Literal['', 'B', 'I', 'BI'], str]] = [
            ('', 'times.ttf'), ('B', 'timesbd.ttf'),
            ('I', 'timesi.ttf'), ('BI', 'timesbi.ttf'),
        ]
        for style, font_file in font_styles:
            path = os.path.join(self.base_fonts_path, font_file)
            if os.path.exists(path):
                pdf.add_font('TimesNewRoman', style, path, uni=True)
            else:
                logger.warning(f"Файл шрифта не найден: {path}")

    def create(self, content: str, filename: str) -> None:
        json_content = self._to_json(content)

        pdf = PDF()
        
        self._setup_fonts(pdf)

        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        for step in json_content['instructions']:
            if pdf.get_y() > 240: 
                pdf.add_page()

            pdf.set_font('TimesNewRoman', 'B', 14)
            pdf.cell(0, 10, step['title'], ln=True)

            pdf.set_font('TimesNewRoman', '', 14)
            pdf.multi_cell(0, 8, step['description'])
            pdf.ln(3)

            # Вставляем картинку
            pdf.image(step['best_image_id'], x=10, w=100)
            pdf.ln(1)

            pdf.set_font('TimesNewRoman', 'I', 10)
            pdf.cell(0, 10, f"Время: {self.format_time(float(step['start_time']))} - {self.format_time(float(step['end_time']))}", ln=True)
            pdf.ln(5)

        pdf.output(filename)

class DocumentFactory:
    _rigestry: Dict[str, Type[BaseDocument]] = {}

    @classmethod
    def register(cls, ext: str, doc_cls: Type[BaseDocument]) -> None:
        cls._rigestry[ext.lower()] = doc_cls
    @classmethod
    def create_document(cls, ext: str) -> BaseDocument:
        ext = ext.lower()
        if ext not in cls._rigestry:
            raise ValueError(f"Неподдерживаемое расширение: {ext}")
        return cls._rigestry[ext]()

DocumentFactory.register(ext=".docx", doc_cls=DocxDocument)
DocumentFactory.register(ext=".pdf", doc_cls=PdfDocument)

class DocumentCreator:
    @staticmethod
    def create(text: str, path: str) -> None:
        _, doc_suffix = os.path.splitext(path)
        document = DocumentFactory.create_document(ext=doc_suffix)
        document.create(text, path)
